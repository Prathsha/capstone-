"""
routers/export.py
-----------------
Word document export endpoint.

POST /api/export/docx
  Body: { content: str, doc_type: str, account_name: str }
  Returns: .docx file as a streaming download.

Parses the Markdown-like output from Gemini (##, ###, -, **, table rows)
and builds a properly structured Word document using python-docx.
"""

import io
import re
from typing import Any

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

router = APIRouter(prefix="/api")


# ── Request model ─────────────────────────────────────────────────────────────

class ExportRequest(BaseModel):
    content: str
    doc_type: str       # "strategy" or "sales_play"
    account_name: str


# ── Filename sanitizer ────────────────────────────────────────────────────────

def _safe_filename(account_name: str, doc_type: str) -> str:
    """Build a safe .docx filename from account name and doc type."""
    safe = re.sub(r"[^\w\s-]", "", account_name)
    safe = re.sub(r"\s+", "_", safe.strip())
    return f"{safe}_{doc_type}.docx"


# ── Markdown → docx renderer ──────────────────────────────────────────────────

def _build_docx(content: str) -> bytes:
    """
    Convert the Gemini Markdown-like reply into a Word document.

    Supported patterns (in order of precedence):
      ## text        → Heading 1
      ### text       → Heading 2
      #### text      → Heading 3
      | col | col |  → Table row (consecutive rows collected into one table)
      - text / • text → Bullet list item
      **text**       → Bold paragraph
      (empty line)   → skip
      anything else  → Normal paragraph
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="python-docx is not installed. Run: pip install python-docx",
        )

    doc = Document()

    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Headings ──────────────────────────────────────────────────────────
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=3)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
            i += 1
            continue

        # ── Table: collect consecutive pipe-rows ──────────────────────────────
        if stripped.startswith("|"):
            table_rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                # Skip separator rows (cells made of only dashes/colons)
                if not all(re.match(r"^[-:]+$", c) for c in cells if c):
                    table_rows.append(cells)
                i += 1

            if table_rows:
                # Normalise column count
                col_count = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=col_count)
                table.style = "Table Grid"
                for r_idx, row_cells in enumerate(table_rows):
                    for c_idx in range(col_count):
                        cell_text = row_cells[c_idx] if c_idx < len(row_cells) else ""
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.text = cell_text
                        # Bold the header row
                        if r_idx == 0:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.bold = True
            continue

        # ── Bullet items ──────────────────────────────────────────────────────
        if stripped.startswith("- ") or stripped.startswith("• "):
            text = stripped[2:]
            para = doc.add_paragraph(style="List Bullet")
            _add_inline_bold(para, text)
            i += 1
            continue

        # ── Bold standalone line (**text**) ───────────────────────────────────
        if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            para = doc.add_paragraph()
            run = para.add_run(stripped[2:-2])
            run.bold = True
            i += 1
            continue

        # ── Empty line ────────────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Normal paragraph (may contain inline **bold**) ────────────────────
        para = doc.add_paragraph()
        _add_inline_bold(para, stripped)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _add_inline_bold(para: Any, text: str) -> None:
    """
    Add runs to *para*, bolding text surrounded by **…**.
    Falls back to a single plain run if no bold markers are present.
    """
    segments = re.split(r"(\*\*[^*]+\*\*)", text)
    for seg in segments:
        if seg.startswith("**") and seg.endswith("**"):
            run = para.add_run(seg[2:-2])
            run.bold = True
        else:
            para.add_run(seg)


# ── Endpoint ──────────────────────────────────────────────────────────────────

@router.post("/export/docx")
async def export_docx(req: ExportRequest):
    """
    Convert a Gemini-generated Markdown reply into a .docx and return it
    as a file download.
    """
    docx_bytes = _build_docx(req.content)
    filename = _safe_filename(req.account_name, req.doc_type)

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
